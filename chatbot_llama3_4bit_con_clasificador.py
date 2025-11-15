
# Chat guiado con plantillas + clasificación en español
# DOCX profesional tipo “Protocolo de instalación” (NFPA 80)

import os, sys, traceback, re, unicodedata
import torch
import torch.nn.functional as F
import numpy as np
import gradio as gr
from datetime import datetime
from transformers import pipeline
from langdetect import detect, LangDetectException  # pip install langdetect
from deep_translator import GoogleTranslator


from transformers import (
    AutoTokenizer, AutoModelForCausalLM,
    BitsAndBytesConfig, BertTokenizer, BertForSequenceClassification,
    RobertaTokenizer, RobertaForSequenceClassification
)
from huggingface_hub import login as hf_login, HfFolder
import joblib

# ---------- DOCX requerido ----------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    print("\n[ERROR] Falta 'python-docx'. Instálalo con:\n   pip install python-docx\n")
    sys.exit(1)

# =============== Control de Modo (ADICIÓN) ===============
# Estado global para modo conversacional o modo reporte
_mode_state = {"mode": "chat"}  # valores posibles: "chat" o "report"

def _detect_intent_for_mode(user_input: str) -> str:
    """
    Detecta si el usuario quiere crear un reporte o cancelar.
    Retorna:
      - "create_report" si detecta deseo de generar documento,
      - "cancel_report" si detecta que el usuario no desea el reporte,
      - "chat" en cualquier otro caso.
    """
    if not isinstance(user_input, str):
        return "chat"
    t = user_input.lower()
    # intención de crear reporte
    if re.search(r"\b(crear|generar|hacer|elaborar).*(reporte|informe|documento)\b", t):
        return "create_report"
    # intención de cancelar reporte o expresar "no quiero"
    if re.search(r"\b(cancelar|no quiero|dejar|salir|no generar|no hagas)\b", t):
        return "cancel_report"
    return "chat"

# =============== Config LLM ===============
MODEL_NAME         = "meta-llama/Meta-Llama-3-8B-Instruct"
MAX_INPUT_TOKENS   = 768
MAX_NEW_TOKENS     = 600
TEMPERATURE        = 0.65
TOP_P              = 0.9
TOP_K              = 40
REPETITION_PENALTY = 1.1

# =============== Token HF ===============

HF_TOKEN = os.environ.get("HF_TOKEN") or HfFolder.get_token()
if HF_TOKEN:
    try:
        hf_login(HF_TOKEN)
        HfFolder.save_token(HF_TOKEN)
        print("Hugging Face: autenticado desde variable de entorno o caché.")
    except Exception as e:
        print("Advertencia HF:", e)
else:
    print("⚠️ No se detectó HF_TOKEN. Ejecuta `huggingface-cli login` o define la variable de entorno HF_TOKEN.")

# --- Traductor español → inglés (versión robusta) ---
translator = pipeline("translation", model="Helsinki-NLP/opus-mt-es-en")

def translate_to_english(text):
    """
    Traduce automáticamente al inglés si detecta idioma español o texto no en inglés.
    Usa langdetect + heurísticas simples.
    """
    if not isinstance(text, str) or not text.strip():
        return text

    t = text.strip()
    try:
        lang = detect(t)
    except LangDetectException:
        lang = None

    # Palabras clave y heurísticas
    spanish_markers = [
        " la ", " el ", " las ", " los ", "puerta", "bisagra",
        "rápid", "demasiado", "cierra", "cerrar", "defecto", "ajuste", "marco"
    ]

    # Condiciones para traducir
    should_translate = (
        lang == "es"
        or (lang is None and any(m in t.lower() for m in spanish_markers))
        or (not all(ord(c) < 128 for c in t))  # caracteres no ASCII
    )

    if should_translate:
        try:
            translation = translator(t)[0]['translation_text']
            print(f"[Translator] '{t}' → '{translation}'")  # ✅ log útil en consola
            return translation
        except Exception as e:
            print(f"[Translator error] {e} ; returning original text")
            return t

    return t

# =============== Utilidades VRAM/DTYPE ===============
def _compute_dtype():
    if torch.cuda.is_available():
        major, _ = torch.cuda.get_device_capability(0)
        return torch.bfloat16 if major >= 8 else torch.float16
    return torch.float32

def _gpu_mem_ok(min_free_gib=6.0):
    if not torch.cuda.is_available():
        return False, "CUDA no disponible"
    try:
        free, total = torch.cuda.mem_get_info()
        free_gib = free / (1024**3)
        return (free_gib >= min_free_gib), f"VRAM libre: {free_gib:.1f} GiB"
    except Exception as e:
        return True, f"No se pudo consultar VRAM: {e}"

def _truncate_input_ids(tokenizer, input_ids, max_len):
    if input_ids.shape[1] <= max_len: return input_ids
    return input_ids[:, -max_len:]

# =============== Tokenizer ===============
print("Cargando tokenizer ...")
try:
    llama_tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, token=HF_TOKEN, use_fast=True)
    if llama_tokenizer.pad_token_id is None:
        llama_tokenizer.pad_token = llama_tokenizer.eos_token
except Exception as e:
    print("Error cargando tokenizer:", e); sys.exit(1)

# =============== Cargar LLM (4-bit → fallback 8-bit offload) ===============
llama_model = None
if torch.cuda.is_available():
    ok, msg = _gpu_mem_ok(6.0)
    print(msg)
    try:
        print("Intentando 4-bit NF4 (todo en GPU) ...")
        bnb_4bit = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_use_double_quant=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_compute_dtype=_compute_dtype(),
        )
        llama_model = AutoModelForCausalLM.from_pretrained(
            MODEL_NAME,
            quantization_config=bnb_4bit,
            device_map={"": 0},
            trust_remote_code=False,
            low_cpu_mem_usage=True,
            use_safetensors=True,
            token=HF_TOKEN,
        )
        llama_model.eval()
        print("LLaMA-3 8B cargado en 4-bit (NF4) 100% GPU.")
    except Exception:
        print("4-bit no entró. Fallback a 8-bit con offload CPU ...")
        try:
            bnb_8bit = BitsAndBytesConfig(load_in_8bit=True, llm_int8_enable_fp32_cpu_offload=True)
            offload_dir = "./offload_llama3_8bit"
            os.makedirs(offload_dir, exist_ok=True)
            llama_model = AutoModelForCausalLM.from_pretrained(
                MODEL_NAME,
                quantization_config=bnb_8bit,
                device_map="auto",
                offload_folder=offload_dir,
                trust_remote_code=False,
                low_cpu_mem_usage=True,
                use_safetensors=True,
                token=HF_TOKEN,
            )
            llama_model.eval()
            print("LLaMA-3 8B cargado en 8-bit con offload GPU+CPU.")
        except Exception:
            print("Error en fallback 8-bit:")
            traceback.print_exc(); sys.exit(1)
else:
    print("No hay CUDA. Este flujo requiere GPU."); sys.exit(1)

# =============== Clasificador (CPU) ===============
DEVICE_CLF = torch.device("cpu")
BERT_CKPT     = "bert-finetuned-epoch7-acc0.9812"
ROBERTA_CKPT  = "roberta-finetuned-epoch6-acc0.9812"
LABEL_ENCODER = "label_encoder.pkl"
MAX_LEN_CLF   = 128
BEST_W        = 0.50

CLASS_NAMES = []

# Mapeo EN -> ES (ajusta a tus clases reales)
LABELS_ES_MAP = {
    "scratch": "rayadura",
    "dent": "abolladura",
    "corrosion": "corrosión",
    "crack": "grieta",
    "discoloration": "decoloración",
    "burr": "rebaba",
    "impact": "impacto",
    "abrasion": "abrasión",
    "handling": "manejo",
    "transport": "transporte",
    "manufacturing": "fabricación",
    "others":"otros",
}

bert = roberta = bert_tok = rob_tok = None

def _to_spanish_label(label_en: str) -> str:
    if not isinstance(label_en, str):
        return str(label_en)
    return LABELS_ES_MAP.get(label_en.lower(), label_en)

def _try_load_classifier():
    global bert, roberta, bert_tok, rob_tok, CLASS_NAMES
    try:
        bert_tok = BertTokenizer.from_pretrained(BERT_CKPT, local_files_only=True)
        bert = BertForSequenceClassification.from_pretrained(BERT_CKPT, local_files_only=True).to(DEVICE_CLF)
        rob_tok = RobertaTokenizer.from_pretrained(ROBERTA_CKPT, local_files_only=True)
        roberta = RobertaForSequenceClassification.from_pretrained(ROBERTA_CKPT, local_files_only=True).to(DEVICE_CLF)
        label_encoder = joblib.load(LABEL_ENCODER)
        CLASS_NAMES = list(label_encoder.classes_)
        print("Clasificador híbrido cargado (local, CPU).")
        return True
    except Exception as e_local:
        print("Local falló:", e_local)
        try:
            bert_tok = BertTokenizer.from_pretrained(BERT_CKPT, token=HF_TOKEN)
            bert = BertForSequenceClassification.from_pretrained(BERT_CKPT, token=HF_TOKEN).to(DEVICE_CLF)
            rob_tok = RobertaTokenizer.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN)
            roberta = RobertaForSequenceClassification.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN).to(DEVICE_CLF)
            label_encoder = joblib.load(LABEL_ENCODER)
            CLASS_NAMES = list(label_encoder.classes_)
            print("Clasificador híbrido cargado (descargado, CPU).")
            return True
        except Exception as e_remote:
            print("Error cargando clasificador:", e_remote)
            return False

_classifier_ok = _try_load_classifier()

# =============== Limpieza de texto para el clasificador ===============
def _clean_for_clf(s: str) -> str:
    if not isinstance(s, str):
        s = str(s) if s is not None else ""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[^\S\r\n]+", " ", s)
    s = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", s)
    s = re.sub(r"([!?¿¡.,;:])\1{2,}", r"\1\1", s)
    return s.strip()

def _translate_if_needed(text: str) -> str:
    try:
        if not text.strip():
            return text
        # Traduce automáticamente de cualquier idioma a inglés
        translated = GoogleTranslator(source='auto', target='en').translate(text)
        if translated and translated.lower() != text.lower():
            print(f"[DEBUG] Texto traducido automáticamente: {translated}")
        return translated
    except Exception as e:
        print(f"[WARN] No se pudo traducir: {e}")
        return text



@torch.no_grad()
def classify_text_hybrid(text: str, best_w: float = BEST_W):
    text = _clean_for_clf(text)
    text = _translate_if_needed(text)  # 🔹 Traducción automática antes de clasificar

    if not _classifier_ok or not (bert and roberta and bert_tok and rob_tok and CLASS_NAMES):
        return {"error": "Clasificador no cargado. Revisa checkpoints y label_encoder.pkl."}
    eb = bert_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)
    er =  rob_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)
    pb = F.softmax(bert(**eb).logits, dim=1)
    pr = F.softmax(roberta(**er).logits, dim=1)
    probs = (best_w * pb + (1 - best_w) * pr)[0].detach().cpu().numpy()
    pred_idx = int(np.argmax(probs))
    label_en = CLASS_NAMES[pred_idx]
    label_es = _to_spanish_label(label_en)
    return {
        "label": label_en,          # interno
        "label_es": label_es,       # mostrado al usuario / docx
        "confidence": float(probs[pred_idx] * 100),
    }

# =============== Checklists SÍ/NO/N/A ===============
CHK_MARCO = [
    ("chk_m_unic_dim", "Unicación y dimensiones de vano terminado es conforme"),
    ("chk_m_nivel_piso", "Nivel de Piso terminado y sin desnivel"),
    ("chk_m_armado_marco", "Armado de piezas de marco de puerta (a nivel de piso) es conforme"),
    ("chk_m_sin_defectos", "No presenta defectos físicos (abolladuras, quiñes, etc.)"),
    ("chk_m_fijacion", "Perforaciones y fijación de marco con tirafones"),
    ("chk_m_verticalidad", "La instalación cumple con los requerimientos de verticalidad y acabado"),
    ("chk_m_otros", "Otros _____"),
]
CHK_HOJA = [
    ("chk_h_aplomado", "Marco de puerta está aplomado a nivel"),
    ("chk_h_medidas_planos", "Medidas nominales de la hoja de puerta están de acuerdo a los Planos del Proyecto"),
    ("chk_h_inst_hoja", "Correcta instalación de hoja de puerta (incluyen bisagras)"),
    ("chk_h_cierre_juego", "Cierre y juego adecuado"),
    ("chk_h_acabado", "Acabado superficial conforme"),
    ("chk_h_limpieza", "Limpieza final del área de trabajo"),
    ("chk_h_otros", "Otros _____"),
]

# =============== Plantillas ===============
TEMPLATES = {
    "informe_incidente": {
        "title": "Informe de Incidente",
        "fields": [
            # Cabecera protocolo
            {"key":"cliente","label":"Cliente", "required":False},
            {"key":"contratista","label":"Contratista", "required":False},
            {"key":"subcontratista","label":"Subcontratista", "required":False},
            {"key":"piso_sector","label":"Piso / Sector", "required":False},
            {"key":"supervision","label":"Supervisión", "required":False},
            {"key":"nro_protocolo","label":"Número de protocolo", "required":False},

            {"key": "fecha", "label": "Fecha del incidente", "required": True, "hint":"Ej: 08/10/2025"},
            {"key": "lugar", "label": "Lugar", "required": True, "hint":"Ej: Planta A - Línea 3"},
            {"key": "ubicacion","label":"Ubicación de la puerta", "required":False},

            {"key": "responsable", "label": "Responsable", "required": True, "hint":"Nombre y cargo"},
            {"key": "acciones", "label": "Acciones correctivas", "required": False, "hint":"Medidas aplicadas y plan"},

            # Checklist Marco (sí/no/n/a)
            *[{"key": k, "label": f"(Marco) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_MARCO],
            # Checklist Hoja
            *[{"key": k, "label": f"(Hoja) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_HOJA],

            {"key": "observaciones", "label": "Observaciones", "required": False, "hint":"Notas a clasificar"},
        ],
    },
}

# =============== Helpers de intents naturales ===============
def _normalize(s: str) -> str:
    s = s.lower().strip()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    return s
YES_SET = {"si","sí","s","x","true","1","yes"}
NO_SET  = {"no","n","false","0"}

def _parse_yesno(value: str):
    """Devuelve 'si' o 'no' si es válido; None si es inválido."""
    if value is None:
        return None
    t = _normalize(str(value))
    if t in YES_SET:
        return "si"
    if t in NO_SET:
        return "no"
    # tolerancias típicas: frases largas que contengan 'si' o 'no' aislado
    if re.search(r"\bsi\b|\bsí\b", t):
        return "si"
    if re.search(r"\bno\b", t):
        return "no"
    return None
def _match_template(user_text_norm: str):
    alias = {
        "informe_incidente": ["informe de incidente", "reporte de incidente", "informe incidente", "incidente", "protocolo de instalación"]
    }
    for key, patterns in alias.items():
        if any(p in user_text_norm for p in patterns):
            return key
    for k, meta in TEMPLATES.items():
        if _normalize(meta["title"]) in user_text_norm or k in user_text_norm:
            return k
    return None

def route_intent(user_text: str, state: dict):
    t = _normalize(user_text)
    if any(w in t for w in ["plantilla", "plantillas", "que puedo hacer", "ayuda", "opciones", "menu"]):
        return {"type":"list_templates", "payload":{}}
    if any(w in t for w in ["cancelar", "anular", "reiniciar", "borrar todo", "empezar de nuevo", "reset"]):
        return {"type":"cancel", "payload":{}}
    if any(w in t for w in ["descarga", "descargar", "genera el archivo", "generar archivo", "exporta", "exportar", "documento", "archivo final"]):
        return {"type":"download", "payload":{}}
    tk = _match_template(t)
    if tk:
        return {"type":"start_template", "payload":{"template_key": tk}}
    if state and state.get("active") and not state.get("finished"):
        return {"type":"answer_field", "payload":{"text": user_text}}
    return {"type":"chitchat", "payload":{"text": user_text}}

# =============== Generación LLM (chat) ===============
@torch.no_grad()
def llama_chat_generate(user_message: str, history_pairs: list | None = None) -> str:
    history_pairs = history_pairs or []
    messages = [{
        "role":"system",
        "content":(
            "Eres un asistente que guía al usuario para crear documentos a partir de plantillas y "
            "clasificar automáticamente la sección 'Observaciones'. Si el usuario parece querer iniciar "
            "un 'Informe de Incidente', sugiere iniciar la plantilla correspondiente, "
            "preguntando el primer campo de forma directa y breve."
        )
    }]
    for u,a in history_pairs:
        messages += [{"role":"user","content":u},{"role":"assistant","content":a}]
    messages.append({"role":"user","content":user_message})

    prompt = llama_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    enc = llama_tokenizer(prompt, return_tensors="pt", truncation=False)
    enc["input_ids"] = _truncate_input_ids(llama_tokenizer, enc["input_ids"], MAX_INPUT_TOKENS)
    if "attention_mask" in enc:
        enc["attention_mask"] = enc["attention_mask"][:, -enc["input_ids"].shape[1]:]

    device = next(llama_model.parameters()).device
    inputs = {k: v.to(device) for k,v in enc.items()}

    out = llama_model.generate(
        **inputs,
        do_sample=True,
        max_new_tokens=MAX_NEW_TOKENS,
        temperature=TEMPERATURE,
        top_p=TOP_P,
        top_k=TOP_K,
        repetition_penalty=REPETITION_PENALTY,
        pad_token_id=llama_tokenizer.pad_token_id,
        eos_token_id=llama_tokenizer.eos_token_id,
        use_cache=True,
    )
    gen_ids = out[0]; input_len = inputs["input_ids"].shape[1]
    new_tokens = gen_ids[input_len:]
    return llama_tokenizer.decode(new_tokens, skip_special_tokens=True).strip()

# =============== Wizard por chat (estado) ===============
FOLLOWUP_QUESTIONS = []
CONF_OK = 75.0
MARGIN_OK = 10.0

def _reset_session():
    return {
        "active": False,
        "template_key": None,
        "fields": [],
        "idx": 0,
        "answers": {},
        "awaiting_followup": False,
        "followup_queue": [],
        "followup_answers": {},
        "rounds": 0,
        "cls_result": None,
        "file_path": None,
        "finished": False,
        "custom_collect": None,
    }

def _render_next_question(state):
    f = state["fields"][state["idx"]]
    if f["key"] in CUSTOM_OTHER_KEYS:
        # Inicializa bloque personalizado para 'Otros' y devuelve el mensaje guiado
        prompt = _start_custom_other_block(state, f) if not state.get("custom_collect") else (
            f"Continuemos con **{f['label']}**. Escribe la descripción o responde **no** para avanzar."
        )
        return prompt
    req = " (obligatorio)" if f.get("required") else ""
    hint = f.get("hint") or ""
    return f"**{f['label']}{req}:**\n_{hint}_"

def _reclassify_with_context(obs_text, extra_dict):
    """
    Combina el texto de observaciones con contexto adicional,
    traduce a inglés si está en español, y luego clasifica.
    """
    enrich = []
    for k, v in extra_dict.items():
        if v:
            enrich.append(f"{k}: {v}")

    # Texto principal de observación
    combined = (obs_text or "").strip()
    if enrich:
        combined += "\n" + "\n".join(f"- {e}" for e in enrich)

    # 🔹 Traducción automática antes de clasificar
    combined_translated = translate_to_english(combined)

    # 🔹 Clasificación híbrida (BERT + RoBERTa)
    cls_result = classify_text_hybrid(combined_translated)

    return cls_result, combined_translated


CUSTOM_OTHER_KEYS = {"chk_m_otros", "chk_h_otros"}

def _start_custom_other_block(state, field):
    prompt = (
        f"En el apartado **{field['label']}** puedes agregar descripciones adicionales. "
        "Escribe el texto del punto extra o responde **no** para omitir."
    )
    state["custom_collect"] = {
        "field_key": field["key"],
        "label": field["label"],
        "entries": [],
        "step": "ask_desc",
        "current_desc": None,
    }
    return prompt

def _handle_custom_collect(state, user_msg):
    data = state.get("custom_collect")
    if not data:
        return None
    text = (user_msg or "").strip()
    if data["step"] == "ask_desc":
        if not text:
            return {"message": "Escribe una descripción o responde **no** para continuar.", "advance": False}
        if _normalize(text) == "no":
            state["answers"][data["field_key"]] = list(data["entries"])
            state["custom_collect"] = None
            message = (
                "No se agregarán más puntos en 'Otros'."
                if data["entries"] else
                "Sin puntos adicionales en 'Otros'."
            )
            return {"message": message, "advance": True}
        data["current_desc"] = text
        data["step"] = "ask_mark"
        return {
            "message": f"¿Marcamos \"{text}\" como **sí** o **no**?",
            "advance": False,
        }
    if data["step"] == "ask_mark":
        choice = _parse_yesno(text)
        if choice is None:
            return {"message": "Responde solo **sí** o **no** para ese punto.", "advance": False}
        data["entries"].append({"text": data["current_desc"], "value": choice})
        data["current_desc"] = None
        data["step"] = "ask_desc"
        return {
            "message": "¿Deseas agregar otro punto en 'Otros'? Escribe la descripción o responde **no**.",
            "advance": False,
        }
    return {"message": "", "advance": False}

def _need_more_questions(res): return False
def _enqueue_followups(): return []

# ============== Helpers de diseño DOCX (estilo protocolo) ==============
LOGO_PATH = "./logo.png"  # cambia la ruta si tu logo está en otro lugar
ACCENT_HEX = "1F3B73"
PALE_ACCENT = "ECF0FF"
LIGHT_BG = "F8FAFF"
ACCENT_RGB = RGBColor(31, 59, 115)
WHITE_RGB = RGBColor(255, 255, 255)

def _apply_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(4)

    def _ensure_style(name, base_type=WD_STYLE_TYPE.PARAGRAPH, base=None, size=12, bold=False):
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, base_type)
            if base:
                style.base_style = doc.styles[base]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = RGBColor(31, 59, 115)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        return style

    _ensure_style("SectionTitle", base="Heading 2", size=13, bold=True)
    muted = _ensure_style("SectionMuted", size=11)
    muted.font.color.rgb = RGBColor(94, 105, 134)
    body = _ensure_style("CardBody", size=10)
    body.font.color.rgb = RGBColor(50, 60, 82)

def _set_cell_shading(cell, fill_hex="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def _cell_p(cell, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def _add_spacer(doc: Document, h=0.12):
    doc.add_paragraph().paragraph_format.space_after = Pt(h*72)

def _section_heading(doc: Document, title: str, subtitle: str = ""):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, PALE_ACCENT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(31, 59, 115)
    if subtitle:
        sub = cell.add_paragraph(subtitle)
        sub.style = "SectionMuted" if "SectionMuted" in doc.styles else doc.styles["Normal"]
    _add_spacer(doc, 0.05)
    return tbl

def _card_block(doc: Document, title: str, body: str):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, LIGHT_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading = cell.paragraphs[0]
    heading.style = "SectionTitle" if "SectionTitle" in doc.styles else doc.styles["Normal"]
    heading.text = title
    body_par = cell.add_paragraph(body or "—")
    body_par.style = "CardBody" if "CardBody" in doc.styles else doc.styles["Normal"]
    _add_spacer(doc, 0.08)
    return tbl

def _signature_block(doc: Document):
    entries = [
        ("Inspector responsable", "________________________    Fecha: ____________"),
        ("Responsable del proyecto", "________________________    Fecha: ____________"),
    ]
    tbl = doc.add_table(rows=len(entries), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, line) in enumerate(entries):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = line
    _add_spacer(doc, 0.1)
    return tbl

def _kv_table(doc, headers_and_values, cols=4):
    rows = (len(headers_and_values) + cols//2) // (cols//2)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_i = c_i = 0
    for k, v in headers_and_values:
        label_cell = tbl.cell(r_i, c_i)
        value_cell = tbl.cell(r_i, c_i+1)
        _set_cell_shading(label_cell, "E3E8F7")
        _set_cell_shading(value_cell, "FFFFFF")
        _cell_p(label_cell, k, bold=True, size=10, color=ACCENT_RGB)
        _cell_p(value_cell, v or "", size=10)
        c_i += 2
        if c_i >= cols:
            c_i = 0
            r_i += 1
    _add_spacer(doc, 0.10)
    return tbl

def _is_yes(v: str) -> bool:
    return _parse_yesno(v) == "si"

def _is_na(v: str) -> bool:
    if v is None: return False
    t = _normalize(str(v))
    return t in {"n/a","na","no aplica","noaplica"}

def _mark_for_box(v: str) -> str:
    """SÍ -> X ; NO u otro -> vacío"""
    return "X" if _is_yes(v) else ""

def _clean_check_label(label: str) -> str:
    return re.sub(r"^\((Marco|Hoja)\)\s*", "", label).strip()

def _expand_checklist_items(items_with_keys, answers):
    rows = []
    for key, label in items_with_keys:
        if key in CUSTOM_OTHER_KEYS:
            extras = answers.get(key, [])
            if not isinstance(extras, list) or not extras:
                continue
            for idx, entry in enumerate(extras, start=1):
                desc = (entry.get("text") or f"{label} #{idx}").strip()
                rows.append((desc, entry.get("value", "")))
        else:
            rows.append((_clean_check_label(label), answers.get(key, "")))
    return rows

def _items_check_table(doc, titulo_seccion, items_with_values):
    _section_heading(doc, titulo_seccion, "Marca con X los puntos conformes")
    t_hdr = doc.add_table(rows=1, cols=2)
    t_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = t_hdr.rows[0].cells
    _set_cell_shading(c0, ACCENT_HEX); _cell_p(c0, "DESCRIPCIÓN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=WHITE_RGB)
    _set_cell_shading(c1, ACCENT_HEX); _cell_p(c1, "SÍ / NO",     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=WHITE_RGB)

    t = doc.add_table(rows=len(items_with_values), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(items_with_values):
        if i % 2 == 0:
            _set_cell_shading(t.cell(i,0), "F4F6FB")
            _set_cell_shading(t.cell(i,1), "F4F6FB")
        _cell_p(t.cell(i,0), label, size=10)
        pc = t.cell(i,1).paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pc.add_run(_mark_for_box(val))
        run.font.size = Pt(12)
    _add_spacer(doc, 0.15)
    return t

# =============== Generación DOCX con diseño (NFPA 80) ===============
def _generate_doc(template_key, answers, cls_res):
    title_img = "PROTOCOLO DE INSTALACIÓN DE PUERTAS CORTAFUEGO"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = os.path.join(os.getcwd(), "outputs"); os.makedirs(out_dir, exist_ok=True)
    fname = "Protocolo_Instalacion_Puertas_" if template_key=="informe_incidente" else TEMPLATES[template_key]["title"]
    fpath = os.path.join(out_dir, f"{fname}_{ts}.docx")

    doc = Document()
    _apply_base_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.5)

    # Encabezado con logo y TÍTULO CENTRADO
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = hdr_tbl.rows[0].cells
    if os.path.exists(LOGO_PATH):
        try:
            left.paragraphs[0].add_run().add_picture(LOGO_PATH, width=Inches(1.2))
        except Exception:
            _cell_p(left, " ", size=1)
    else:
        _cell_p(left, " ", size=1)
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run(title_img)
    r.bold = True; r.font.size = Pt(13)
    _add_spacer(doc, 0.05)

    if template_key == "informe_incidente":
        # Cabecera protocolo
        _section_heading(doc, "Datos generales del protocolo", "Identificación del cliente y obra")
        _kv_table(
            doc,
            [
                ("CLIENTE:",        str(answers.get("cliente",""))),
                ("FECHA:",          str(answers.get("fecha",""))),
                ("CONTRATISTA:",    str(answers.get("contratista",""))),
                ("UBICACIÓN DE LA PUERTA:", str(answers.get("ubicacion",""))),
                ("SUBCONTRATISTA:", str(answers.get("subcontratista",""))),
                ("PISO/SECTOR:",    str(answers.get("piso_sector",""))),
                ("SUPERVISIÓN:",    str(answers.get("supervision",""))),
                ("NÚMERO DE PROTOCOLO:", str(answers.get("nro_protocolo",""))),
            ],
            cols=4
        )

        resumen = []
        if answers.get("lugar"):
            resumen.append(f"Lugar inspeccionado: {answers.get('lugar')}")
        if answers.get("responsable"):
            resumen.append(f"Responsable: {answers.get('responsable')}")
        if resumen:
            _card_block(doc, "Resumen del incidente", "\n".join(resumen))

        # Puntos de control “Marco”
        marco_vals = _expand_checklist_items(CHK_MARCO, answers)
        _items_check_table(doc, "Puntos de control - Marco", marco_vals)

        # Puntos de control “Hoja”
        hoja_vals = _expand_checklist_items(CHK_HOJA, answers)
        _items_check_table(doc, "Puntos de control - Hoja", hoja_vals)

        # Observaciones generales
        obs_text = str(answers.get("observaciones","")).strip()
        if obs_text:
            _card_block(doc, "Observaciones reportadas", obs_text)

        # Comentarios / Acciones
        acciones_text = str(answers.get("acciones","")).strip() or "No se registraron acciones correctivas."
        _card_block(doc, "Comentarios y acciones correctivas", acciones_text)

        # Clasificación automática (una etiqueta en español)
        if answers.get("observaciones",""):
            if cls_res and "error" not in cls_res:
                label = cls_res.get("label_es", cls_res.get("label",""))
                cls_body = f"Etiqueta predicha: {label}\nConfianza: {cls_res['confidence']:.2f}%"
                _card_block(doc, "Clasificación automática de observaciones", cls_body)
            else:
                _card_block(doc, "Clasificación automática de observaciones", "Clasificador no disponible.")

        # Firmas
        _section_heading(doc, "Firmas de conformidad", "Validación de responsables")
        _signature_block(doc)

    doc.save(fpath)
    return fpath

# =============== Chat Handler con enrutamiento natural (MODIFICADO para modo) ===============
def chat_handler(message, chat_history, session_state, file_comp, quick_action=None):
    chat_history = chat_history or []
    state = session_state or _reset_session()
    user_msg = (message or "").strip()
    if quick_action:
        user_msg = quick_action

    pairs, last_user = [], None
    for m in chat_history:
        if m.get("role") == "user": last_user = m.get("content")
        elif m.get("role") == "assistant" and last_user is not None:
            pairs.append((last_user, m.get("content"))); last_user = None

    def send(bot_text):
        if user_msg:
            chat_history.append({"role":"user","content":user_msg})
        chat_history.append({"role":"assistant","content":bot_text})
        return chat_history

    # Intent routing basado en el contenido y el estado de la plantilla
    intent = route_intent(user_msg, state)

    # —————— DETECCIÓN Y SWITCH DE MODO (ADICIÓN) ——————
    # Detectamos intención de modo (crear/cancelar/rellenar)
    intent_mode = _detect_intent_for_mode(user_msg)
    if intent_mode == "create_report":
        _mode_state["mode"] = "report"
    elif intent_mode == "cancel_report":
        # Si el usuario explícitamente cancela, volvemos a chat y reiniciamos sesión activa
        _mode_state["mode"] = "chat"
        # además, si había una plantilla activa, la reiniciamos
        if state and state.get("active"):
            state = _reset_session()

        reply = "Entendido, salimos del modo de reporte. Puedes seguir conversando libremente conmigo."
        chat_history.append({"role":"user","content":user_msg})
        chat_history.append({"role":"assistant","content":reply})
        return chat_history, "", state, gr.update()

    # Si estamos en modo 'chat' y no hay intención explícita de plantilla/listado/descarga/cancel,
    # generamos respuesta libre y devolvemos (evitamos forzar flujo plantilla).
    if _mode_state.get("mode") == "chat":
        # si el intent apunta a start_template/list_templates/download/cancel, dejamos que el flujo original lo maneje
        if intent["type"] not in {"start_template", "list_templates", "download", "cancel"} and not (state and state.get("active") and not state.get("finished")):
            # Modo charla libre → llamar al LLM conversacional
            try:
                reply = llama_chat_generate(user_msg, pairs)
            except Exception as e:
                reply = f"Error generando respuesta: {e}"
            if user_msg:
                chat_history.append({"role":"user","content":user_msg})
            chat_history.append({"role":"assistant","content":reply})
            return chat_history, "", state, gr.update()

    # —————— Intents globales (mantengo exactamente tu lógica previa) ——————
    if intent["type"] == "list_templates":
        names = "\n".join([f"- **{meta['title']}** (di: *Quiero {meta['title']}*)" for _,meta in TEMPLATES.items()])
        tips = ("Puedo iniciar una plantilla, preguntar los **checklists (sí/no/n/a)** "
                "y generar el Word con casilleros marcados automáticamente.")
        return send(f"{tips}\n\nPlantillas disponibles:\n{names}"), "", state, gr.update()

    if intent["type"] == "cancel":
        state = _reset_session()
        return send("He cancelado la sesión actual. ¿Deseas iniciar una nueva plantilla? (Ej.: *Quiero un Informe de Incidente*)"), "", state, gr.update()

    if intent["type"] == "download":
        if not state["active"] or not state["finished"]:
            return send("Aún no has completado la plantilla. Si ya terminaste de responder, dime *descargar* cuando estés listo."), "", state, gr.update()
        fpath = _generate_doc(state["template_key"], state["answers"], state["cls_result"])
        state["file_path"] = fpath
        return send(f"Documento generado: **{os.path.basename(fpath)}** (ver debajo para descargar)"), "", state, gr.update(value=fpath)

    if intent["type"] == "start_template":
        key = intent["payload"]["template_key"]
        state = _reset_session()
        state["active"] = True
        state["template_key"] = key
        state["fields"] = TEMPLATES[key]["fields"]
        state["idx"] = 0
        q = _render_next_question(state)
        # When we explicitly start a template, ensure mode switches to report
        _mode_state["mode"] = "report"
        return send(f"Perfecto, iniciaré **{TEMPLATES[key]['title']}**.\n{q}"), "", state, gr.update(value=None)

    # Flujo guiado
    if (intent["type"] == "answer_field") and state["active"] and not state["finished"]:

        if state["awaiting_followup"] and state["followup_queue"]:
            key, _ = state["followup_queue"].pop(0)
            state["followup_answers"][key] = user_msg
            if state["followup_queue"]:
                next_q = state["followup_queue"][0][1]
                return send(f"Gracias. {next_q}"), "", state, gr.update()
            obs = state["answers"].get("observaciones", "")
            res, combined = _reclassify_with_context(obs, state["followup_answers"])
            state["answers"]["observaciones"] = combined
            state["cls_result"] = res
            state["rounds"] += 1
            if "error" in res:
                state["awaiting_followup"] = False
                state["finished"] = True
                return send(f"Clasificador no disponible: {res['error']}\nHas completado la plantilla. Di **descargar** para generar el documento."), "", state, gr.update()
            msg = f"**Clasificación**: {res.get('label_es', res['label'])} ({res['confidence']:.2f}%)"
            state["awaiting_followup"] = False
            state["finished"] = True
            return send(f"{msg}\n\n¡Listo! Has completado la plantilla. Di **descargar** para obtener el archivo."), "", state, gr.update()

        f = state["fields"][state["idx"]]

        if state.get("custom_collect"):
            custom = _handle_custom_collect(state, user_msg)
            if not custom or not custom["advance"]:
                reply = custom["message"] if custom and custom.get("message") else "Necesito una respuesta válida para 'Otros'."
                return send(reply), "", state, gr.update()
            # Avanzar al siguiente campo tras cerrar el bloque de 'Otros'
            state["idx"] += 1
            if state["idx"] < len(state["fields"]):
                q = _render_next_question(state)
                base_msg = custom.get("message") or "Continuemos."
                combo = f"{base_msg}\n\n{q}" if base_msg else q
                return send(combo), "", state, gr.update()
            state["finished"] = True
            closing = custom.get("message") or "¡Plantilla completa! Di **descargar** para generar el documento."
            return send(closing), "", state, gr.update()

        if f.get("required") and not user_msg:
            return send(f"El campo **{f['label']}** es obligatorio. Intenta nuevamente."), "", state, gr.update()
        if f["key"].startswith("chk_"):
            yn = _parse_yesno(user_msg)
            if yn is None:
                return send(
                    f"⚠️ Opción no válida para **{f['label']}**. Usa solo **sí** o **no** "
                    "(también acepto: si/s/x/yes/1 y no/n/0/false). Intenta de nuevo:"
                ), "", state, gr.update()
            state["answers"][f["key"]] = yn
        else:
            state["answers"][f["key"]] = user_msg

        just_obs = (f["key"] == "observaciones")
        state["idx"] += 1

        if just_obs and user_msg.strip():
            res = classify_text_hybrid(user_msg)
            state["cls_result"] = res
            if "error" in res:
                state["finished"] = True
                return send(f"Clasificador no disponible: {res['error']}\nHas completado la plantilla. Di **descargar** para generar el documento."), "", state, gr.update()
            # mostramos solo 1 etiqueta en español
            _ = f"**Clasificación**: {res.get('label_es', res['label'])} ({res['confidence']:.2f}%)"

        if state["idx"] < len(state["fields"]):
            q = _render_next_question(state)
            return send(q), "", state, gr.update()

        state["finished"] = True
        return send("¡Plantilla completa! Di **descargar** para generar el documento."), "", state, gr.update()

    # Si llegamos aquí y no entró en ninguno de los flujos anteriores,
    # (p. ej. casos raros), intentamos una respuesta de fallback de charla.
    try:
        reply = llama_chat_generate(user_msg, pairs)
    except Exception as e:
        reply = f"Error generando respuesta: {e}"
    if user_msg:
        chat_history.append({"role":"user","content":user_msg})
    chat_history.append({"role":"assistant","content":reply})
    return chat_history, "", state, gr.update()

# =============== UI Gradio (dark + moderno) ===============
DARK_CSS = """
<style>
  :root { --radius-lg: 16px; --blur-bg: rgba(12, 18, 32, 0.78); }
  body {
    background: #050910;
    color: #e6e9ef;
    font-family: 'Inter', ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, "Helvetica Neue", Arial, "Noto Sans";
    position: relative;
    overflow-x: hidden;
  }
  body::before {
    content: "";
    position: fixed;
    inset: 0;
    background: radial-gradient(circle at 20% 20%, rgba(79,111,255,0.15), transparent 40%),
                radial-gradient(circle at 80% 0%, rgba(147,70,255,0.15), transparent 45%),
                radial-gradient(circle at 50% 80%, rgba(23,188,255,0.12), transparent 40%);
    filter: blur(40px);
    animation: aurora 24s ease-in-out infinite alternate;
    pointer-events: none;
    z-index: -2;
  }
  .gradio-container {
    max-width: 1220px !important;
    margin: 0 auto !important;
    padding: 32px 28px 64px !important;
    background: radial-gradient(circle at top, rgba(49,91,255,0.18), transparent 55%) #050910;
  }
  .block, .group, .gr-panel, .gr-box, .form, .tabs, .tabitem, .tabs > div, .gradio-row {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
  }
  .hero-card {
    background: var(--blur-bg);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 24px;
    padding: 22px 26px;
    margin-bottom: 14px;
    backdrop-filter: blur(16px);
    box-shadow: 0 18px 38px rgba(0,0,0,0.45);
    opacity: 0;
    animation: floatUp 0.8s ease forwards;
  }
  .hero-card h1 {
    font-size: 28px;
    margin: 0;
    color: #f5f7ff;
  }
  .hero-sub {
    color: rgba(230,233,239,0.85);
    margin-top: 4px;
  }
  .hero-pill {
    display:inline-flex;
    align-items:center;
    gap:8px;
    padding:8px 14px;
    border-radius:999px;
    border:1px solid rgba(255,255,255,0.12);
    background: rgba(13,30,60,0.55);
    font-size:13px;
    animation: pulseGlow 3s ease-in-out infinite;
  }
  .main-layout {
    gap: 22px;
    align-items: stretch;
  }
  .panel {
    background: var(--blur-bg);
    border: 1px solid rgba(255,255,255,0.06);
    border-radius: 24px;
    padding: 18px 20px 22px;
    box-shadow: 0 10px 35px rgba(0,0,0,0.4);
    opacity: 0;
    animation: fadeIn 0.85s ease forwards;
  }
  .chat-panel .gradio-chatbot {
    background: rgba(7,12,22,0.9) !important;
    border: 1px solid rgba(255,255,255,0.04) !important;
    border-radius: 20px !important;
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.02), 0 15px 45px rgba(2,6,16,0.8);
    animation: floatDelay 1s ease forwards;
  }
  .chat-panel .gradio-chatbot > div {
    background: transparent !important;
  }
  .message.user, .chatbot .message.user {
    background: #111a29 !important;
    border-radius: 14px !important;
    border: 1px solid rgba(255,255,255,0.04) !important;
    animation: chatPop 0.35s ease;
  }
  .message.bot, .chatbot .message.bot {
    background: #0c131f !important;
    border-radius: 14px !important;
    border: 1px solid rgba(21,38,63,0.8) !important;
    animation: chatPop 0.35s ease;
  }
  .chat-panel .gr-textbox textarea {
    min-height: 54px !important;
  }
  .chat-panel .gr-textbox, textarea, .gr-input, input[type="text"] {
    background: rgba(8,13,23,0.9) !important;
    color: #f1f4ff !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
    border-radius: 999px !important;
    padding: 14px 18px !important;
    font-size: 16px !important;
  }
  .chat-panel .chips-row {
    flex-wrap: wrap;
    gap: 10px;
    margin-bottom: 4px;
  }
  .chip-button {
    flex: 1 1 48%;
    min-width: 180px;
    background: rgba(25,47,87,0.7) !important;
    border: 1px solid rgba(65,110,255,0.35) !important;
    border-radius: 14px !important;
    color: #f0f3ff !important;
    font-weight: 600 !important;
    height: 48px !important;
    transition: transform .2s ease, box-shadow .2s ease, border .2s ease, filter .2s ease;
    animation: cascade 0.9s ease forwards;
  }
  .chip-button:nth-child(2n) { animation-delay: 0.15s; }
  .chip-button:nth-child(3n) { animation-delay: 0.3s; }
  .chip-button:nth-child(4n) { animation-delay: 0.45s; }
  .chip-button:hover {
    transform: translateY(-2px) scale(1.01);
    box-shadow: 0 12px 24px rgba(31,63,122,0.35);
    filter: brightness(1.05);
  }
  .file-card {
    margin-top: 12px;
    border-radius: 20px !important;
    background: rgba(10,17,30,0.9) !important;
    border: 1px dashed rgba(90,119,255,0.45) !important;
  }
  .info-panel .info-card {
    background: rgba(7,12,22,0.75);
    border: 1px solid rgba(255,255,255,0.05);
    border-radius: 18px;
    padding: 18px 20px;
    margin-bottom: 14px;
    opacity: 0;
    animation: slideIn 0.85s ease forwards;
  }
  .info-panel h3 {
    margin-top: 0;
    font-size: 18px;
    color: #f5f7ff;
  }
  .info-steps ol {
    padding-left: 22px;
    margin: 0;
    color: rgba(230,233,239,0.85);
  }
  .info-tags {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
  }
  .info-tags span {
    padding: 6px 12px;
    border-radius: 999px;
    border: 1px solid rgba(255,255,255,0.08);
    background: rgba(12,20,34,0.8);
    font-size: 13px;
  }
  .status-board {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
    gap: 10px;
  }
  .status-item {
    background: rgba(10,16,28,0.85);
    border: 1px solid rgba(255,255,255,0.05);
    border-radius: 16px;
    padding: 12px 14px;
    font-size: 13px;
    animation: cardPulse 2.4s ease-in-out infinite;
  }
  .status-item span {
    display: block;
    color: rgba(230,233,239,0.68);
    font-size: 12px;
  }
  @media (max-width: 980px) {
    .main-layout { flex-direction: column; }
    .chip-button { flex: 1 1 100%; }
  }
  @keyframes aurora {
    0% { transform: translate3d(-4%, -2%, 0) scale(1); }
    100% { transform: translate3d(6%, 4%, 0) scale(1.05); }
  }
  @keyframes floatUp {
    0% { opacity: 0; transform: translateY(20px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes fadeIn {
    0% { opacity: 0; transform: translateY(24px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes floatDelay {
    0% { opacity: 0; transform: translateY(18px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes slideIn {
    0% { opacity: 0; transform: translateX(30px); }
    100% { opacity: 1; transform: translateX(0); }
  }
  @keyframes cascade {
    0% { opacity: 0; transform: translateY(20px) scale(0.98); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
  }
  @keyframes pulseGlow {
    0%, 100% { box-shadow: 0 0 10px rgba(98,130,255,0.25); }
    50% { box-shadow: 0 0 18px rgba(98,130,255,0.45); }
  }
  @keyframes chatPop {
    0% { opacity: 0; transform: translateY(8px) scale(0.98); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
  }
  @keyframes cardPulse {
    0%, 100% { border-color: rgba(255,255,255,0.05); }
    50% { border-color: rgba(126,162,255,0.18); }
  }
</style>
"""

def build_and_launch_gradio():
    with gr.Blocks(title="Chat Plantillas + Clasificación (NFPA 80)", css=DARK_CSS, theme=gr.themes.Soft(
        primary_hue="indigo", neutral_hue="slate"
    )) as demo:

        gr.HTML("""
        <div class="hero-card">
          <div style="display:flex;justify-content:space-between;align-items:flex-start;gap:18px;flex-wrap:wrap;">
            <div>
              <h1>Asistente Inteligente NFPA 80</h1>
              <p class="hero-sub">Completa plantillas guiadas, obtiene una clasificación automática<br/>y genera documentos Word listos para entregar.</p>
            </div>
            <div class="hero-pill" id="status-pill">
              <span>Modo actual</span>
              <strong>Listo para empezar</strong>
            </div>
          </div>
        </div>
        """)

        with gr.Row(elem_classes=["main-layout"]):
            with gr.Column(scale=7, elem_classes=["panel", "chat-panel"]):
                chatbot = gr.Chatbot(height=520, type="messages", label=None)
                txt = gr.Textbox(
                    show_label=False,
                    placeholder="Haz una pregunta o di “Quiero un Informe de Incidente”…",
                    lines=1,
                    max_lines=1,
                    autofocus=True,
                )
                with gr.Row(elem_classes=["chips-row"]):
                    qa_create_inf = gr.Button("📝 Informe de incidente", elem_classes=["chip-button"])
                with gr.Row(elem_classes=["chips-row"]):
                    qa_list = gr.Button("📚 Ver plantillas", elem_classes=["chip-button"])
                    qa_download = gr.Button("⬇️ Descargar documento", elem_classes=["chip-button"])
                file_out = gr.File(label="Documento generado", interactive=False, elem_classes=["file-card"])
            with gr.Column(scale=5, elem_classes=["panel", "info-panel"]):
                gr.HTML("""
                <div class="info-card info-steps">
                  <h3>Cómo funciona</h3>
                  <ol>
                    <li>Inicia una plantilla con los botones rápidos o escribiendo “Quiero…”.</li>
                    <li>Responde las preguntas guiadas y los checklists sí/no.</li>
                    <li>El clasificador detecta la categoría de la observación.</li>
                    <li>Descarga el DOCX formateado con protocolo NFPA 80.</li>
                  </ol>
                </div>
                """)
                gr.HTML("""
                <div class="info-card">
                  <h3>Sugerencias rápidas</h3>
                  <div class="info-tags">
                    <span>Modo chat libre para dudas</span>
                    <span>Automatiza checklists</span>
                    <span>Traducción al inglés integrada</span>
                    <span>Clasificación híbrida BERT + RoBERTa</span>
                  </div>
                </div>
                """)
                state = gr.State(_reset_session())

        txt.submit(chat_handler, [txt, chatbot, state, file_out], [chatbot, txt, state, file_out])
        qa_create_inf.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Quiero informe de incidente", visible=False), chatbot, state, file_out, gr.Textbox(value="Quiero informe de incidente", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )
        qa_list.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Ver plantillas", visible=False), chatbot, state, file_out, gr.Textbox(value="Ver plantillas", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )
        qa_download.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Descargar documento", visible=False), chatbot, state, file_out, gr.Textbox(value="Descargar documento", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )

    demo.launch(server_name="127.0.0.1", server_port=7860, share=False)

# =============== MAIN ===============
if __name__ == "__main__":
    try:
        build_and_launch_gradio()
    except Exception as e:
        print("Error lanzando Gradio:", e)
        traceback.print_exc(); sys.exit(1)
