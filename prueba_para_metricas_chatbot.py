#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Chat guiado con plantillas + clasificación en español
DOCX profesional tipo “Protocolo de instalación” (NFPA 80)
- Correcciones de sintaxis y saltos de línea
- Manejo robusto de HF_TOKEN
- Carga LLaMA-3 en 4-bit (NF4) con fallback 8-bit offload
- Traducción unificada (langdetect + GoogleTranslator)
- Clasificador híbrido BERT+RoBERTa (CPU) con top-3 probabilidades
- Preguntas de seguimiento automáticas si la confianza es baja
- UI Gradio modernizada (chips de acciones rápidas)
- DOCX: incluye etiqueta predicha y top-3
"""

import os
import sys
import traceback
import re
import unicodedata
from datetime import datetime
from typing import List, Tuple, Dict, Any, Optional

import numpy as np
import torch
import torch.nn.functional as F
import gradio as gr

from langdetect import detect, LangDetectException  # pip install langdetect
from deep_translator import GoogleTranslator        # pip install deep-translator

from transformers import (
    AutoTokenizer,
    AutoModelForCausalLM,
    BitsAndBytesConfig,
    BertTokenizer,
    BertForSequenceClassification,
    RobertaTokenizer,
    RobertaForSequenceClassification,
)
from huggingface_hub import login as hf_login, HfFolder
import joblib

# ---------- DOCX requerido ----------
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    print("\n[ERROR] Falta 'python-docx'. Instálalo con:\n   pip install python-docx\n")
    sys.exit(1)

# =============== Config LLM ===============
MODEL_NAME         = "meta-llama/Meta-Llama-3-8B-Instruct"
MAX_INPUT_TOKENS   = 768
MAX_NEW_TOKENS     = 256
TEMPERATURE        = 0.7
TOP_P              = 0.9
TOP_K              = 40
REPETITION_PENALTY = 1.1

# =============== Token HF ===============
# Usa la variable de entorno HF_TOKEN. No hardcodees el token.
HF_TOKEN = (os.environ.get("HF_TOKEN") or "").strip()
if HF_TOKEN:
    try:
        hf_login(HF_TOKEN)
        HfFolder.save_token(HF_TOKEN)
        print("Hugging Face: autenticado.")
    except Exception as e:
        print("[WARN] No se pudo autenticar en HF:", e)
else:
    print("[INFO] No se detectó HF_TOKEN. Si el repo es gated, fallará la descarga.")

# --- Traducción (es/auto → en) ---
# Heurística: detecta idioma; si es español u otro distinto a inglés, traduce a inglés.

def translate_to_english(text: str) -> str:
    if not isinstance(text, str) or not text.strip():
        return text
    t = text.strip()
    try:
        lang = detect(t)
    except LangDetectException:
        lang = None

    spanish_markers = [
        " la ", " el ", " las ", " los ", "puerta", "bisagra",
        "rápid", "demasiado", "cierra", "cerrar", "defecto", "ajuste", "marco"
    ]

    should_translate = (
        lang == "es"
        or (lang is None and any(m in t.lower() for m in spanish_markers))
        or (lang not in (None, "en"))
        or (not all(ord(c) < 128 for c in t))
    )

    if should_translate:
        try:
            translation = GoogleTranslator(source='auto', target='en').translate(t)
            print(f"[Translator] '{t[:120]}' → '{translation[:120]}'")
            return translation
        except Exception as e:
            print(f"[Translator error] {e}; devolviendo texto original")
            return t
    return t

# =============== Utilidades VRAM/DTYPE ===============

def _compute_dtype():
    if torch.cuda.is_available():
        major, _ = torch.cuda.get_device_capability(0)
        return torch.bfloat16 if major >= 8 else torch.float16
    return torch.float32


def _gpu_mem_ok(min_free_gib=6.0) -> Tuple[bool, str]:
    if not torch.cuda.is_available():
        return False, "CUDA no disponible"
    try:
        free, _total = torch.cuda.mem_get_info()
        free_gib = free / (1024**3)
        return (free_gib >= min_free_gib), f"VRAM libre: {free_gib:.1f} GiB"
    except Exception as e:
        return True, f"No se pudo consultar VRAM: {e}"


def _truncate_input_ids(tokenizer, input_ids, max_len):
    if input_ids.shape[1] <= max_len:
        return input_ids
    return input_ids[:, -max_len:]

# =============== Tokenizer ===============
print("Cargando tokenizer ...")
try:
    llama_tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, token=HF_TOKEN or None, use_fast=True)
    if llama_tokenizer.pad_token_id is None:
        llama_tokenizer.pad_token = llama_tokenizer.eos_token
except Exception as e:
    print("Error cargando tokenizer:", e)
    sys.exit(1)

# =============== Cargar LLM (4-bit → fallback 8-bit offload) ===============
llama_model = None
if torch.cuda.is_available():
    ok, msg = _gpu_mem_ok(6.0)
    print(msg)
    try:
        print("Intentando 4-bit NF4 (todo en GPU) ...")
        bnb_4bit = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_use_double_quant=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_compute_dtype=_compute_dtype(),
        )
        llama_model = AutoModelForCausalLM.from_pretrained(
            MODEL_NAME,
            quantization_config=bnb_4bit,
            device_map={"": 0},
            trust_remote_code=False,
            low_cpu_mem_usage=True,
            use_safetensors=True,
            token=HF_TOKEN or None,
        )
        llama_model.eval()
        print("LLaMA-3 8B cargado en 4-bit (NF4) 100% GPU.")
    except Exception:
        print("[WARN] 4-bit no entró. Fallback a 8-bit con offload CPU ...")
        try:
            bnb_8bit = BitsAndBytesConfig(load_in_8bit=True, llm_int8_enable_fp32_cpu_offload=True)
            offload_dir = "./offload_llama3_8bit"
            os.makedirs(offload_dir, exist_ok=True)
            llama_model = AutoModelForCausalLM.from_pretrained(
                MODEL_NAME,
                quantization_config=bnb_8bit,
                device_map="auto",
                offload_folder=offload_dir,
                trust_remote_code=False,
                low_cpu_mem_usage=True,
                use_safetensors=True,
                token=HF_TOKEN or None,
            )
            llama_model.eval()
            print("LLaMA-3 8B cargado en 8-bit con offload GPU+CPU.")
        except Exception:
            print("Error en fallback 8-bit:")
            traceback.print_exc()
            sys.exit(1)
else:
    print("No hay CUDA. Este flujo requiere GPU.")
    sys.exit(1)

# =============== Clasificador (CPU) ===============
DEVICE_CLF = torch.device("cpu")
BERT_CKPT     = "bert-finetuned-epoch7-acc0.9812"
ROBERTA_CKPT  = "roberta-finetuned-epoch6-acc0.9812"
LABEL_ENCODER = "label_encoder.pkl"
MAX_LEN_CLF   = 128
BEST_W        = 0.50

CLASS_NAMES: List[str] = []

# Mapeo EN -> ES (ajústalo a tus clases reales)
LABELS_ES_MAP = {
    "scratch": "rayadura",
    "dent": "abolladura",
    "corrosion": "corrosión",
    "crack": "grieta",
    "discoloration": "decoloración",
    "burr": "rebaba",
    "impact": "impacto",
    "abrasion": "abrasión",
    "handling": "manejo",
    "transport": "transporte",
    "manufacturing": "fabricación",
    "others": "otros",
}

bert = roberta = bert_tok = rob_tok = None


def _to_spanish_label(label_en: str) -> str:
    if not isinstance(label_en, str):
        return str(label_en)
    return LABELS_ES_MAP.get(label_en.lower(), label_en)


def _try_load_classifier() -> bool:
    global bert, roberta, bert_tok, rob_tok, CLASS_NAMES
    try:
        bert_tok = BertTokenizer.from_pretrained(BERT_CKPT, local_files_only=True)
        bert = BertForSequenceClassification.from_pretrained(BERT_CKPT, local_files_only=True).to(DEVICE_CLF)
        rob_tok = RobertaTokenizer.from_pretrained(ROBERTA_CKPT, local_files_only=True)
        roberta = RobertaForSequenceClassification.from_pretrained(ROBERTA_CKPT, local_files_only=True).to(DEVICE_CLF)
        label_encoder = joblib.load(LABEL_ENCODER)
        CLASS_NAMES = list(label_encoder.classes_)
        print("Clasificador híbrido cargado (local, CPU).")
        return True
    except Exception as e_local:
        print("[WARN] Local falló:", e_local)
        try:
            bert_tok = BertTokenizer.from_pretrained(BERT_CKPT, token=HF_TOKEN or None)
            bert = BertForSequenceClassification.from_pretrained(BERT_CKPT, token=HF_TOKEN or None).to(DEVICE_CLF)
            rob_tok = RobertaTokenizer.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN or None)
            roberta = RobertaForSequenceClassification.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN or None).to(DEVICE_CLF)
            label_encoder = joblib.load(LABEL_ENCODER)
            CLASS_NAMES = list(label_encoder.classes_)
            print("Clasificador híbrido cargado (descargado, CPU).")
            return True
        except Exception as e_remote:
            print("[ERROR] Cargando clasificador:", e_remote)
            return False


_classifier_ok = _try_load_classifier()

# =============== Limpieza de texto para el clasificador ===============

def _clean_for_clf(s: str) -> str:
    if not isinstance(s, str):
        s = str(s) if s is not None else ""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[^\S\r\n]+", " ", s)
    s = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", s)
    s = re.sub(r"([!?¿¡.,;:])\1{2,}", r"\1\1", s)
    return s.strip()


def _translate_if_needed(text: str) -> str:
    try:
        if not text.strip():
            return text
        translated = GoogleTranslator(source='auto', target='en').translate(text)
        if translated and translated.lower() != text.lower():
            print(f"[DEBUG] Texto traducido automáticamente: {translated[:120]}")
        return translated
    except Exception as e:
        print(f"[WARN] No se pudo traducir: {e}")
        return text


@torch.no_grad()
def classify_text_hybrid(text: str, best_w: float = BEST_W) -> Dict[str, Any]:
    text = _clean_for_clf(text)
    text = _translate_if_needed(text)  # Traducción automática antes de clasificar

    if not _classifier_ok or not (bert and roberta and bert_tok and rob_tok and CLASS_NAMES):
        return {"error": "Clasificador no cargado. Revisa checkpoints y label_encoder.pkl."}

    eb = bert_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)
    er = rob_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)

    pb = F.softmax(bert(**eb).logits, dim=1)
    pr = F.softmax(roberta(**er).logits, dim=1)
    probs = (best_w * pb + (1 - best_w) * pr)[0].detach().cpu().numpy()

    top_idx = int(np.argmax(probs))
    top_label_en = CLASS_NAMES[top_idx]
    top_label_es = _to_spanish_label(top_label_en)

    # top-3
    order = np.argsort(probs)[::-1][:3]
    top3 = [
        {
            "label": CLASS_NAMES[i],
            "label_es": _to_spanish_label(CLASS_NAMES[i]),
            "confidence": float(probs[i] * 100),
        } for i in order
    ]

    return {
        "label": top_label_en,
        "label_es": top_label_es,
        "confidence": float(probs[top_idx] * 100),
        "top3": top3,
    }

# =============== Checklists SÍ/NO/N/A ===============
CHK_MARCO = [
    ("chk_m_unic_dim", "Ubicación y dimensiones de vano terminado es conforme"),
    ("chk_m_nivel_piso", "Nivel de Piso terminado y sin desnivel"),
    ("chk_m_armado_marco", "Armado de piezas de marco de puerta (a nivel de piso) es conforme"),
    ("chk_m_sin_defectos", "No presenta defectos físicos (abolladuras, quiñes, etc.)"),
    ("chk_m_fijacion", "Perforaciones y fijación de marco con tirafones"),
    ("chk_m_verticalidad", "La instalación cumple con los requerimientos de verticalidad y acabado"),
    ("chk_m_otros", "Otros _____"),
]

CHK_HOJA = [
    ("chk_h_aplomado", "Marco de puerta está aplomado a nivel"),
    ("chk_h_medidas_planos", "Medidas nominales de la hoja de puerta están de acuerdo a los Planos del Proyecto"),
    ("chk_h_inst_hoja", "Correcta instalación de hoja de puerta (incluyen bisagras)"),
    ("chk_h_cierre_juego", "Cierre y juego adecuado"),
    ("chk_h_acabado", "Acabado superficial conforme"),
    ("chk_h_limpieza", "Limpieza final del área de trabajo"),
    ("chk_h_otros", "Otros _____"),
]

# =============== Plantillas ===============
TEMPLATES: Dict[str, Dict[str, Any]] = {
    "informe_incidente": {
        "title": "Informe de Incidente",
        "fields": [
            {"key": "cliente", "label": "Cliente", "required": False},
            {"key": "contratista", "label": "Contratista", "required": False},
            {"key": "plano_referencia", "label": "Plano de referencia", "required": False},
            {"key": "subcontratista", "label": "Subcontratista", "required": False},
            {"key": "piso_sector", "label": "Piso / Sector", "required": False},
            {"key": "supervision", "label": "Supervisión", "required": False},
            {"key": "nro_protocolo", "label": "Número de protocolo", "required": False},

            {"key": "fecha", "label": "Fecha del incidente", "required": True, "hint": "Ej: 08/10/2025"},
            {"key": "lugar", "label": "Lugar", "required": True, "hint": "Ej: Planta A - Línea 3"},
            {"key": "tipo_puerta", "label": "Tipo de puerta", "required": False},
            {"key": "ubicacion", "label": "Ubicación de la puerta", "required": False},

            {"key": "responsable", "label": "Responsable", "required": True, "hint": "Nombre y cargo"},
            {"key": "descripcion", "label": "Descripción", "required": True, "hint": "¿Qué ocurrió? Detalles clave"},
            {"key": "causa", "label": "Causa probable", "required": False, "hint": "Posible causa raíz"},
            {"key": "acciones", "label": "Acciones correctivas", "required": False, "hint": "Medidas aplicadas y plan"},

            *[{"key": k, "label": f"(Marco) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_MARCO],
            *[{"key": k, "label": f"(Hoja) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_HOJA],

            {"key": "observaciones", "label": "Observaciones", "required": False, "hint": "Notas a clasificar"},
        ],
    },
    "carta_solicitud": {
        "title": "Carta de Solicitud",
        "fields": [
            {"key": "fecha", "label": "Fecha", "required": True},
            {"key": "destinatario", "label": "Destinatario", "required": True},
            {"key": "remitente", "label": "Remitente", "required": True},
            {"key": "asunto", "label": "Asunto", "required": True},
            {"key": "cuerpo", "label": "Cuerpo", "required": True},
            {"key": "despedida", "label": "Despedida", "required": False},
            {"key": "firma", "label": "Firma", "required": False},
        ],
    },
}

# =============== Helpers de intents naturales ===============

def _normalize(s: str) -> str:
    s = s.lower().strip()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    return s

YES_SET = {"si", "sí", "s", "x", "true", "1", "yes"}
NO_SET  = {"no", "n", "false", "0"}


def _parse_yesno(value: Optional[str]) -> Optional[str]:
    """Devuelve 'si' o 'no' si es válido; None si es inválido."""
    if value is None:
        return None
    t = _normalize(str(value))
    if t in YES_SET:
        return "si"
    if t in NO_SET:
        return "no"
    if re.search(r"\bsi\b|\bsí\b", t):
        return "si"
    if re.search(r"\bno\b", t):
        return "no"
    return None


def _match_template(user_text_norm: str) -> Optional[str]:
    alias = {
        "informe_incidente": [
            "informe de incidente", "reporte de incidente", "informe incidente", "incidente", "protocolo de instalación"
        ],
        "carta_solicitud": ["carta de solicitud", "carta", "solicitud"],
    }
    for key, patterns in alias.items():
        if any(p in user_text_norm for p in patterns):
            return key
    for k, meta in TEMPLATES.items():
        if _normalize(meta["title"]) in user_text_norm or k in user_text_norm:
            return k
    return None


def route_intent(user_text: str, state: Dict[str, Any]) -> Dict[str, Any]:
    t = _normalize(user_text)
    if any(w in t for w in ["plantilla", "plantillas", "que puedo hacer", "ayuda", "opciones", "menu"]):
        return {"type": "list_templates", "payload": {}}
    if any(w in t for w in ["cancelar", "anular", "reiniciar", "borrar todo", "empezar de nuevo", "reset"]):
        return {"type": "cancel", "payload": {}}
    if any(w in t for w in ["descarga", "descargar", "genera el archivo", "generar archivo", "exporta", "exportar", "documento", "archivo final"]):
        return {"type": "download", "payload": {}}
    tk = _match_template(t)
    if tk:
        return {"type": "start_template", "payload": {"template_key": tk}}
    if state and state.get("active") and not state.get("finished"):
        return {"type": "answer_field", "payload": {"text": user_text}}
    return {"type": "chitchat", "payload": {"text": user_text}}

# =============== Generación LLM (chat) ===============

@torch.no_grad()
def llama_chat_generate(user_message: str, history_pairs: Optional[List[Tuple[str, str]]] = None) -> str:
    history_pairs = history_pairs or []
    messages = [{
        "role": "system",
        "content": (
            "Eres un asistente que guía al usuario para crear documentos a partir de plantillas y "
            "clasificar automáticamente la sección 'Observaciones'. Si el usuario parece querer iniciar "
            "un 'Informe de Incidente' o una 'Carta de Solicitud', sugiere iniciar la plantilla correspondiente, "
            "preguntando el primer campo de forma directa y breve."
        )
    }]
    for u, a in history_pairs:
        messages.append({"role": "user", "content": u})
        messages.append({"role": "assistant", "content": a})
    messages.append({"role": "user", "content": user_message})

    prompt = llama_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    enc = llama_tokenizer(prompt, return_tensors="pt", truncation=False)
    enc["input_ids"] = _truncate_input_ids(llama_tokenizer, enc["input_ids"], MAX_INPUT_TOKENS)
    if "attention_mask" in enc:
        enc["attention_mask"] = enc["attention_mask"][:, -enc["input_ids"].shape[1]:]

    device = next(llama_model.parameters()).device
    inputs = {k: v.to(device) for k, v in enc.items()}

    out = llama_model.generate(
        **inputs,
        do_sample=True,
        max_new_tokens=MAX_NEW_TOKENS,
        temperature=TEMPERATURE,
        top_p=TOP_P,
        top_k=TOP_K,
        repetition_penalty=REPETITION_PENALTY,
        pad_token_id=llama_tokenizer.pad_token_id,
        eos_token_id=llama_tokenizer.eos_token_id,
        use_cache=True,
    )
    gen_ids = out[0]
    input_len = inputs["input_ids"].shape[1]
    new_tokens = gen_ids[input_len:]
    return llama_tokenizer.decode(new_tokens, skip_special_tokens=True).strip()

# =============== Wizard por chat (estado) ===============
CONF_OK = 75.0           # confianza para aceptar sin preguntas
MARGIN_OK = 10.0         # margen para top-1 vs top-2

FOLLOWUP_POOL = {
    # Preguntas sugeridas por tipo de defecto (ajusta a tu dominio)
    "scratch": [
        "¿La marca afecta la superficie pintada o el material base?",
        "¿La longitud aproximada de la rayadura supera 5 cm?",
    ],
    "dent": [
        "¿Se observa hundimiento del panel o deformación visible?",
        "¿El golpe afecta al cierre o al alineamiento?",
    ],
    "corrosion": [
        "¿Hay presencia de óxido o pérdida de material?",
        "¿Se ubica en zonas cercanas a bisagras o al marco?",
    ],
    "crack": [
        "¿La grieta atraviesa el panel o es superficial?",
        "¿Se extiende desde los bordes o perforaciones?",
    ],
}


def _need_followups(res: Dict[str, Any]) -> bool:
    if not res or "error" in res:
        return False
    conf = float(res.get("confidence", 0.0))
    top3 = res.get("top3", [])
    if len(top3) >= 2:
        gap = top3[0]["confidence"] - top3[1]["confidence"]
    else:
        gap = 100.0
    return (conf < CONF_OK) or (gap < MARGIN_OK)


def _make_followups(res: Dict[str, Any]) -> List[Tuple[str, str]]:
    """Devuelve una lista [(key, question), ...] para solicitar más contexto."""
    if not res or "error" in res:
        return []
    label = str(res.get("label", "others")).lower()
    qlist = FOLLOWUP_POOL.get(label, [
        "¿El defecto afecta el cierre o la hermeticidad de la puerta?",
        "¿En qué zona exacta se ubica (bisagras, borde, panel, marco)?",
    ])
    return [(f"fup_{i+1}", q) for i, q in enumerate(qlist[:2])]


def _reset_session() -> Dict[str, Any]:
    return {
        "active": False,
        "template_key": None,
        "fields": [],
        "idx": 0,
        "answers": {},
        "awaiting_followup": False,
        "followup_queue": [],
        "followup_answers": {},
        "rounds": 0,
        "cls_result": None,
        "file_path": None,
        "finished": False,
    }


def _render_next_question(state: Dict[str, Any]) -> str:
    f = state["fields"][state["idx"]]
    req = " (obligatorio)" if f.get("required") else ""
    hint = f.get("hint") or ""
    return f"**{f['label']}{req}:**\n_{hint}_"


def _reclassify_with_context(obs_text: str, extra_dict: Dict[str, str]):
    enrich = []
    for k, v in extra_dict.items():
        if v:
            enrich.append(f"{k}: {v}")

    combined = (obs_text or "").strip()
    if enrich:
        combined += "\n" + "\n".join(f"- {e}" for e in enrich)

    combined_translated = translate_to_english(combined)
    cls_result = classify_text_hybrid(combined_translated)
    return cls_result, combined_translated

# ============== Helpers de diseño DOCX (estilo protocolo) ==============
LOGO_PATH = "./logo.png"  # cambia la ruta si tu logo está en otro lugar


def _apply_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _set_cell_shading(cell, fill_hex="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _cell_p(cell, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _add_spacer(doc: Document, h=0.12):
    doc.add_paragraph().paragraph_format.space_after = Pt(h * 72)


def _title_bar(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _set_cell_shading(c, "D9D9D9")
    _cell_p(c, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _add_spacer(doc, 0.10)


def _kv_table(doc, headers_and_values, cols=4):
    rows = (len(headers_and_values) + cols // 2) // (cols // 2)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_i = c_i = 0
    for k, v in headers_and_values:
        _cell_p(tbl.cell(r_i, c_i), k, bold=True, size=10)
        _cell_p(tbl.cell(r_i, c_i + 1), v or "", size=10)
        c_i += 2
        if c_i >= cols:
            c_i = 0
            r_i += 1
    _add_spacer(doc, 0.10)
    return tbl


def _spec_table(doc, tipo_puerta="", ubicacion=""):
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_shading(tbl.cell(0, 0))
    _cell_p(tbl.cell(0, 0), "Tipo de Puerta", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shading(tbl.cell(0, 1))
    _cell_p(tbl.cell(0, 1), "Ubicación", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_p(tbl.cell(1, 0), tipo_puerta, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_p(tbl.cell(1, 1), ubicacion, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_spacer(doc, 0.05)
    return tbl


def _is_yes(v: str) -> bool:
    return _parse_yesno(v) == "si"


def _mark_for_box(v: str) -> str:
    return "X" if _is_yes(v) else ""


def _items_check_table(doc, titulo_seccion, items_with_values):
    _title_bar(doc, titulo_seccion)
    t_hdr = doc.add_table(rows=1, cols=2)
    t_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = t_hdr.rows[0].cells
    _set_cell_shading(c0)
    _cell_p(c0, "ITEMS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _set_cell_shading(c1)
    _cell_p(c1, "C", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    t = doc.add_table(rows=len(items_with_values), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(items_with_values):
        _cell_p(t.cell(i, 0), label, size=10)
        pc = t.cell(i, 1).paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pc.add_run(_mark_for_box(val))
        run.font.size = Pt(12)
    _add_spacer(doc, 0.15)
    return t

# =============== Generación DOCX con diseño (NFPA 80) ===============

def _generate_doc(template_key: str, answers: Dict[str, Any], cls_res: Optional[Dict[str, Any]]):
    title_img = "PROTOCOLO DE INSTALACIÓN DE PUERTAS CORTAFUEGO"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = os.path.join(os.getcwd(), "outputs")
    os.makedirs(out_dir, exist_ok=True)

    fname = "Protocolo_Instalacion_Puertas_" if template_key == "informe_incidente" else TEMPLATES[template_key]["title"]
    fpath = os.path.join(out_dir, f"{fname}{ts}.docx")

    doc = Document()
    _apply_base_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.5)

    # Encabezado con logo y título
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = hdr_tbl.rows[0].cells
    if os.path.exists(LOGO_PATH):
        try:
            left.paragraphs[0].add_run().add_picture(LOGO_PATH, width=Inches(1.2))
        except Exception:
            _cell_p(left, " ", size=1)
    else:
        _cell_p(left, " ", size=1)
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run(title_img)
    r.bold = True
    r.font.size = Pt(13)
    _add_spacer(doc, 0.05)

    if template_key == "informe_incidente":
        _kv_table(
            doc,
            [
                ("CLIENTE:",        str(answers.get("cliente", ""))),
                ("FECHA:",          str(answers.get("fecha", ""))),
                ("CONTRATISTA:",    str(answers.get("contratista", ""))),
                ("PLANO DE REFERENCIA:", str(answers.get("plano_referencia", ""))),
                ("SUBCONTRATISTA:", str(answers.get("subcontratista", ""))),
                ("PISO/SECTOR:",    str(answers.get("piso_sector", ""))),
                ("SUPERVISIÓN:",    str(answers.get("supervision", ""))),
                ("NÚMERO DE PROTOCOLO:", str(answers.get("nro_protocolo", ""))),
            ],
            cols=4,
        )

        _title_bar(doc, "ESPECIFICACIONES TÉCNICAS DE LAS PUERTAS CORTAFUEGO")
        _spec_table(doc, tipo_puerta=str(answers.get("tipo_puerta", "")), ubicacion=str(answers.get("ubicacion", "")))

        marco_vals = [(lbl, answers.get(k, "")) for k, lbl in CHK_MARCO]
        _items_check_table(doc, "PUNTO DE CONTROL PARA LA INSTALACIÓN DE MARCO DE PUERTA", marco_vals)

        hoja_vals = [(lbl, answers.get(k, "")) for k, lbl in CHK_HOJA]
        _items_check_table(doc, "PUNTO DE CONTROL PARA LA INSTALACIÓN DE HOJA DE PUERTA", hoja_vals)

        doc.add_heading("Comentarios y acciones correctivas", level=2)
        doc.add_paragraph(str(answers.get("acciones", "")) or "—")
        _add_spacer(doc, 0.1)

        if answers.get("observaciones", ""):
            doc.add_heading("Clasificación automática de Observaciones", level=2)
            if cls_res and "error" not in cls_res:
                # Etiqueta principal
                p = doc.add_paragraph()
                run = p.add_run(
                    f"Etiqueta predicha: {cls_res.get('label_es', cls_res.get('label', ''))} ({cls_res['confidence']:.2f}%)"
                )
                run.bold = True
                # Top-3
                top3 = cls_res.get("top3", [])
                if top3:
                    doc.add_paragraph("Top-3 etiquetas (confianza):")
                    for item in top3:
                        doc.add_paragraph(
                            f"- {item['label_es']} ({item['label']}) — {item['confidence']:.2f}%",
                            style=None,
                        )
            else:
                doc.add_paragraph("Clasificador no disponible.")
            _add_spacer(doc, 0.1)

        doc.add_heading("Firmas", level=2)
        doc.add_paragraph("Inspector: __________________________    Fecha: ____________")
        doc.add_paragraph("Responsable del proyecto: ___________    Fecha: ____________")

    elif template_key == "carta_solicitud":
        doc.add_paragraph(str(answers.get("fecha", "")), style=None)
        _add_spacer(doc, 0.05)
        doc.add_paragraph(str(answers.get("destinatario", "")))
        _add_spacer(doc, 0.15)
        subj = doc.add_paragraph()
        r = subj.add_run("Asunto: ")
        r.bold = True
        subj.add_run(str(answers.get("asunto", "")))
        _add_spacer(doc, 0.1)
        cuerpo = str(answers.get("cuerpo", ""))
        for par in cuerpo.split("\n"):
            if par.strip():
                doc.add_paragraph(par.strip())
        _add_spacer(doc, 0.15)
        doc.add_paragraph(str(answers.get("despedida", "Atentamente,")))
        _add_spacer(doc, 0.15)
        firma = str(answers.get("firma", ""))
        if firma:
            doc.add_paragraph(firma)

    doc.save(fpath)
    return fpath

# =============== Chat Handler con enrutamiento natural ===============

def chat_handler(message, chat_history, session_state, file_comp, quick_action=None):
    chat_history = chat_history or []
    state = session_state or _reset_session()
    user_msg = (message or "").strip()
    if quick_action:
        user_msg = quick_action

    # Construye pares para el contexto del LLM
    pairs, last_user = [], None
    for m in chat_history:
        if m.get("role") == "user":
            last_user = m.get("content")
        elif m.get("role") == "assistant" and last_user is not None:
            pairs.append((last_user, m.get("content", "")))
            last_user = None

    def send(bot_text):
        if user_msg:
            chat_history.append({"role": "user", "content": user_msg})
        chat_history.append({"role": "assistant", "content": bot_text})
        return chat_history

    intent = route_intent(user_msg, state)

    # Intents globales
    if intent["type"] == "list_templates":
        names = "\n".join([f"- **{meta['title']}** (di: *Quiero {meta['title']}*)" for _, meta in TEMPLATES.items()])
        tips = (
            "Puedo iniciar una plantilla, preguntar los **checklists (sí/no)** "
            "y generar el Word con casilleros marcados automáticamente."
        )
        return send(f"{tips}\n\nPlantillas disponibles:\n{names}"), "", state, gr.update()

    if intent["type"] == "cancel":
        state = _reset_session()
        return send("He cancelado la sesión actual. ¿Deseas iniciar una nueva plantilla? (Ej.: *Quiero un Informe de Incidente*)"), "", state, gr.update()

    if intent["type"] == "download":
        if not state["active"] or not state["finished"]:
            return send("Aún no has completado la plantilla. Si ya terminaste de responder, dime *descargar* cuando estés listo."), "", state, gr.update()
        fpath = _generate_doc(state["template_key"], state["answers"], state.get("cls_result"))
        state["file_path"] = fpath
        return send(f"Documento generado: **{os.path.basename(fpath)}** (ver debajo para descargar)"), "", state, gr.update(value=fpath)

    if intent["type"] == "start_template":
        key = intent["payload"]["template_key"]
        state = _reset_session()
        state["active"] = True
        state["template_key"] = key
        state["fields"] = TEMPLATES[key]["fields"]
        state["idx"] = 0
        q = _render_next_question(state)
        return send(f"Perfecto, iniciaré **{TEMPLATES[key]['title']}**.\n{q}"), "", state, gr.update(value=None)

    # Flujo guiado de llenado
    if (intent["type"] == "answer_field") and state["active"] and not state["finished"]:
        f = state["fields"][state["idx"]]
        if f.get("required") and not user_msg:
            return send(f"El campo **{f['label']}** es obligatorio. Intenta nuevamente."), "", state, gr.update()

        if f["key"].startswith("chk_"):
            yn = _parse_yesno(user_msg)
            if yn is None:
                return send(
                    "⚠️ Opción no válida. Usa **sí** o **no** (acepto si/s/x/yes/1 y no/n/0/false). Intenta de nuevo:"
                ), "", state, gr.update()
            state["answers"][f["key"]] = yn
        else:
            state["answers"][f["key"]] = user_msg

        just_obs = (f["key"] == "observaciones")
        state["idx"] += 1

        if just_obs and user_msg.strip():
            # Predicción inicial
            res = classify_text_hybrid(user_msg)
            state["cls_result"] = res

            if "error" in res:
                state["finished"] = True
                return send(f"Clasificador no disponible: {res['error']}\nHas completado la plantilla. Di **descargar** para generar el documento."), "", state, gr.update()

            # ¿Necesitamos preguntas de seguimiento?
            if _need_followups(res):
                state["awaiting_followup"] = True
                state["followup_answers"] = {}
                state["followup_queue"] = _make_followups(res)
                if state["followup_queue"]:
                    _, qf = state["followup_queue"][0]
                    return send(f"Para afinar la clasificación: {qf}"), "", state, gr.update()

        # Manejo de followups
        if state["awaiting_followup"] and state["followup_queue"]:
            key_f, _ = state["followup_queue"].pop(0)
            state["followup_answers"][key_f] = user_msg
            if state["followup_queue"]:
                _next_key, next_q = state["followup_queue"][0]
                return send(f"Gracias. {next_q}"), "", state, gr.update()
            # Re-clasifica con contexto adicional
            obs = state["answers"].get("observaciones", "")
            res2, combined = _reclassify_with_context(obs, state["followup_answers"])
            state["answers"]["observaciones"] = combined
            state["cls_result"] = res2
            state["awaiting_followup"] = False

            if "error" in res2:
                state["finished"] = True
                return send(f"Clasificador no disponible: {res2['error']}\nHas completado la plantilla. Di **descargar** para generar el documento."), "", state, gr.update()

            msg = f"**Clasificación**: {res2.get('label_es', res2['label'])} ({res2['confidence']:.2f}%)"
            state["finished"] = True
            return send(f"{msg}\n\n¡Listo! Has completado la plantilla. Di **descargar** para obtener el archivo."), "", state, gr.update()

        if state["idx"] < len(state["fields"]):
            q = _render_next_question(state)
            return send(q), "", state, gr.update()

        state["finished"] = True
        return send("¡Plantilla completa! Di **descargar** para generar el documento."), "", state, gr.update()

    # Modo charla general
    try:
        reply = llama_chat_generate(user_msg, pairs)
    except Exception as e:
        reply = f"Error generando respuesta: {e}"
    if user_msg:
        chat_history.append({"role": "user", "content": user_msg})
    chat_history.append({"role": "assistant", "content": reply})
    return chat_history, "", state, gr.update()

# =============== UI Gradio (dark + moderno) ===============
DARK_CSS = """
<style>
  :root { --radius-lg: 14px; }
  .gradio-container { background: #0b0f16 !important; color: #e6e9ef !important; font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, "Helvetica Neue", Arial, "Noto Sans"; }
  .block, .group, .gr-panel, .gr-box, .form, .tabs, .tabitem, .tabs > div, .gradio-row { background: transparent !important; border: none !important; }
  .gradio-chatbot, .wrap.svelte-1clx7j5 { background: #0f1623 !important; border: 1px solid #1b2638 !important; border-radius: var(--radius-lg) !important; box-shadow: 0 0 0 1px rgba(255,255,255,0.02) inset, 0 8px 24px rgba(0,0,0,0.35); }
  .message.user, .chatbot .message.user { background: #111a29 !important; border-radius: 12px !important; }
  .message.bot, .chatbot .message.bot { background: #0c131f !important; border-radius: 12px !important; border: 1px solid #162033 !important; }
  .gr-textbox, textarea, .gr-input, input[type="text"] { background: #0f1623 !important; color: #e6e9ef !important; border: 1px solid #1b2638 !important; border-radius: 12px !important; }
  .gr-button, button { background: #122036 !important; color: #e6e9ef !important; border: 1px solid #22314a !important; border-radius: 999px !important; }
  .gr-button:hover { filter: brightness(1.12); }
  .chip { display:inline-block; padding:8px 12px; margin:4px 6px 0 0; border-radius:999px; background:#101a2b; border:1px solid #1e2b45; cursor:pointer; user-select:none; }
  .chip:hover { background:#122036; }
  .header-card { background:#0f1623; border:1px solid #1b2638; border-radius:16px; padding:14px 16px; margin-bottom:8px; }
</style>
"""


def build_and_launch_gradio():
    with gr.Blocks(title="Chat Plantillas + Clasificación (NFPA 80)", css=DARK_CSS, theme=gr.themes.Soft(primary_hue="indigo", neutral_hue="slate")) as demo:
        gr.HTML(
            """
            <div class="header-card">
              <div style="display:flex;justify-content:space-between;align-items:center;gap:12px;flex-wrap:wrap;">
                <div>
                  <div style="font-size:20px;font-weight:700;color:#e6e9ef;">Asistente de Plantillas</div>
                  <div style="opacity:.8">Completa documentos guiados y genera Word (.docx) con diseño de Protocolo NFPA 80.</div>
                </div>
                <div class="chip">Listo para empezar</div>
              </div>
            </div>
            """
        )

        chatbot = gr.Chatbot(height=520, type="messages", label=None)
        txt = gr.Textbox(show_label=False, placeholder="Escribe y pulsa Enter", lines=1, max_lines=1, autofocus=True)

        with gr.Row():
            qa_create_inf = gr.Button("📝 Crear Informe de Incidente", elem_classes=["chip"])
            qa_create_cart = gr.Button("✉️ Carta de Solicitud", elem_classes=["chip"])
            qa_list = gr.Button("📚 Ver plantillas", elem_classes=["chip"])
            qa_download = gr.Button("⬇️ Descargar documento", elem_classes=["chip"])

        file_out = gr.File(label="Descargar documento", interactive=False)
        state = gr.State(_reset_session())

        try:
            dev = next(llama_model.parameters()).device
            dtype = next(llama_model.parameters()).dtype
            gr.Markdown(f"**LLM Device:** `{dev}` &nbsp;&nbsp; **dtype:** `{dtype}` &nbsp;&nbsp; **CUDA:** `{torch.cuda.is_available()}`")
        except Exception:
            gr.Markdown(f"**CUDA:** `{torch.cuda.is_available()}`")

        txt.submit(chat_handler, [txt, chatbot, state, file_out], [chatbot, txt, state, file_out])

        # Clicks: pasamos message="" y quick_action con el texto deseado
        qa_create_inf.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="", visible=False), chatbot, state, file_out, gr.Textbox(value="Quiero informe de incidente", visible=False)],
            outputs=[chatbot, txt, state, file_out],
        )
        qa_create_cart.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="", visible=False), chatbot, state, file_out, gr.Textbox(value="Quiero carta de solicitud", visible=False)],
            outputs=[chatbot, txt, state, file_out],
        )
        qa_list.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="", visible=False), chatbot, state, file_out, gr.Textbox(value="Ver plantillas", visible=False)],
            outputs=[chatbot, txt, state, file_out],
        )
        qa_download.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="", visible=False), chatbot, state, file_out, gr.Textbox(value="Descargar documento", visible=False)],
            outputs=[chatbot, txt, state, file_out],
        )

    demo.launch(server_name="127.0.0.1", server_port=7860, share=False)


# =============== MAIN ===============
if __name__ == "__main__":
    try:
        build_and_launch_gradio()
    except Exception as e:
        print("Error lanzando Gradio:", e)
        traceback.print_exc()
        sys.exit(1)
